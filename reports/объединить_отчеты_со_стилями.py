#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Скрипт для объединения .docx файлов с сохранением исходных стилей и форматирования
"""

import os
import glob
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import parse_xml

def get_docx_files():
    """Получить список всех .docx файлов в текущей директории"""
    docx_files = []
    
    # Ищем все файлы .docx
    for file_path in glob.glob("*.docx"):
        # Пропускаем файлы, которые могут быть результатом объединения
        if not file_path.startswith("объединенный_отчет") and not file_path.startswith("combined_report"):
            docx_files.append(file_path)
    
    # Сортируем файлы по имени для удобства
    docx_files.sort()
    
    return docx_files

def copy_paragraph_with_formatting(source_para, target_doc):
    """Копировать параграф с сохранением форматирования"""
    new_para = target_doc.add_paragraph()
    
    # Копируем выравнивание
    if source_para.alignment:
        new_para.alignment = source_para.alignment
    
    # Копируем стиль
    if source_para.style:
        new_para.style = source_para.style.name
    
    # Копируем отступы
    if source_para.paragraph_format:
        new_para.paragraph_format.left_indent = source_para.paragraph_format.left_indent
        new_para.paragraph_format.right_indent = source_para.paragraph_format.right_indent
        new_para.paragraph_format.first_line_indent = source_para.paragraph_format.first_line_indent
        new_para.paragraph_format.space_before = source_para.paragraph_format.space_before
        new_para.paragraph_format.space_after = source_para.paragraph_format.space_after
        new_para.paragraph_format.line_spacing = source_para.paragraph_format.line_spacing
    
    # Копируем все runs (фрагменты текста с форматированием)
    for run in source_para.runs:
        new_run = new_para.add_run(run.text)
        
        # Копируем свойства шрифта
        if run.font:
            new_run.font.name = run.font.name
            new_run.font.size = run.font.size
            new_run.font.bold = run.font.bold
            new_run.font.italic = run.font.italic
            new_run.font.underline = run.font.underline
            new_run.font.color.rgb = run.font.color.rgb
    
    return new_para

def copy_table_with_formatting(source_table, target_doc):
    """Копировать таблицу с сохранением форматирования"""
    # Создаем таблицу с теми же размерами
    new_table = target_doc.add_table(
        rows=len(source_table.rows),
        cols=len(source_table.columns)
    )
    
    # Копируем стиль таблицы
    if source_table.style:
        new_table.style = source_table.style.name
    
    # Копируем содержимое ячеек
    for i, row in enumerate(source_table.rows):
        for j, cell in enumerate(row.cells):
            new_cell = new_table.cell(i, j)
            
            # Копируем текст из ячейки
            for para in cell.paragraphs:
                copy_paragraph_with_formatting(para, new_cell)
    
    return new_table

def merge_documents_with_styles(docx_files):
    """Объединить документы с сохранением стилей"""
    if not docx_files:
        print("Не найдено .docx файлов для объединения")
        return
    
    print(f"Найдено {len(docx_files)} .docx файлов:")
    for i, file_path in enumerate(docx_files, 1):
        print(f"  {i}. {file_path}")
    
    # Создаем новый документ
    combined_doc = Document()
    
    # Настраиваем поля страницы
    sections = combined_doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Добавляем титульную страницу с простым форматированием
    title = combined_doc.add_heading('ОБЪЕДИНЕННЫЙ ОТЧЕТ', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = combined_doc.add_heading('по всем лабораторным работам', 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    info_para = combined_doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info_para.add_run(f"Объединено {len(docx_files)} отчетов")
    info_run.bold = True
    
    combined_doc.add_page_break()
    
    # Обрабатываем каждый файл
    for i, file_path in enumerate(docx_files, 1):
        print(f"Обработка файла {i}/{len(docx_files)}: {file_path}")
        
        try:
            # Открываем исходный документ
            source_doc = Document(file_path)
            
            # Добавляем заголовок для этого отчета
            report_title = combined_doc.add_heading(f"ОТЧЕТ {i}: {os.path.basename(file_path).replace('.docx', '')}", 1)
            
            # Копируем все параграфы из исходного документа
            for para in source_doc.paragraphs:
                if para.text.strip():  # Пропускаем пустые параграфы
                    copy_paragraph_with_formatting(para, combined_doc)
            
            # Копируем все таблицы из исходного документа
            for table in source_doc.tables:
                copy_table_with_formatting(table, combined_doc)
            
            # Добавляем разрыв страницы между отчетами (кроме последнего)
            if i < len(docx_files):
                combined_doc.add_page_break()
                
        except Exception as e:
            print(f"  Ошибка при обработке файла {file_path}: {e}")
            # Добавляем сообщение об ошибке в документ
            error_para = combined_doc.add_paragraph(f"[Ошибка при обработке файла {file_path}]")
            error_run = error_para.add_run(" - файл не удалось обработать")
            error_run.italic = True
    
    # Сохраняем объединенный документ
    output_filename = "объединенный_отчет_со_стилями.docx"
    combined_doc.save(output_filename)
    
    print(f"\nОбъединенный отчет сохранен в файл: {output_filename}")
    print(f"Размер файла: {os.path.getsize(output_filename)} байт")
    
    return output_filename

def create_simple_version(docx_files):
    """Создать простую версию без сложного форматирования"""
    if not docx_files:
        return
    
    print("\nСоздаю простую версию...")
    
    combined_doc = Document()
    
    # Простой заголовок
    title = combined_doc.add_heading('Объединенный отчет', 0)
    
    # Обрабатываем каждый файл
    for i, file_path in enumerate(docx_files, 1):
        print(f"  Обработка файла {i}/{len(docx_files)}")
        
        try:
            source_doc = Document(file_path)
            
            # Добавляем название файла
            combined_doc.add_heading(os.path.basename(file_path).replace('.docx', ''), 2)
            
            # Копируем только текст (без форматирования)
            for para in source_doc.paragraphs:
                if para.text.strip():
                    combined_doc.add_paragraph(para.text)
            
            # Добавляем разрыв страницы
            if i < len(docx_files):
                combined_doc.add_page_break()
                
        except Exception as e:
            print(f"  Ошибка: {e}")
    
    output_filename = "объединенный_отчет_простой.docx"
    combined_doc.save(output_filename)
    
    print(f"Простая версия сохранена в файл: {output_filename}")
    
    return output_filename

def main():
    """Основная функция"""
    print("=" * 60)
    print("Скрипт для объединения .docx отчетов с сохранением стилей")
    print("=" * 60)
    
    # Получаем список .docx файлов
    docx_files = get_docx_files()
    
    if not docx_files:
        print("В текущей директории не найдено .docx файлов")
        return
    
    # Создаем объединенный документ с сохранением стилей
    output_file = merge_documents_with_styles(docx_files)
    
    # Создаем простую версию
    simple_file = create_simple_version(docx_files)
    
    if output_file:
        print("\nГотово! Объединенные отчеты созданы успешно.")
        print(f"1. Документ со стилями: {os.path.abspath(output_file)}")
        print(f"2. Простой документ: {os.path.abspath(simple_file)}")
        print("\nРекомендуется использовать первый вариант для сохранения исходного форматирования.")

if __name__ == "__main__":
    main()