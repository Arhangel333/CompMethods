#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Скрипт для объединения 4 выбранных .docx файлов с сохранением форматирования и изображений
"""

import os
import zipfile
import shutil
import tempfile
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import parse_xml

def get_selected_files():
    """Получить список выбранных файлов"""
    return [
        "Kurnosov_lab1-2_rep.docx",
        "Kurnosov_lab3-4_rep.docx", 
        "Kurnosov_lab5_report.docx",
        "Kurnosov_lab6_report.docx"
    ]

def extract_images_from_docx(docx_path, temp_dir):
    """Извлечь изображения из .docx файла"""
    images = []
    
    try:
        # Открываем .docx как zip-архив
        with zipfile.ZipFile(docx_path, 'r') as zip_ref:
            # Ищем файлы в папке word/media/
            for file_info in zip_ref.infolist():
                if file_info.filename.startswith('word/media/'):
                    # Извлекаем изображение
                    image_name = os.path.basename(file_info.filename)
                    image_path = os.path.join(temp_dir, image_name)
                    
                    # Извлекаем файл
                    with zip_ref.open(file_info) as source, open(image_path, 'wb') as target:
                        shutil.copyfileobj(source, target)
                    
                    images.append({
                        'name': image_name,
                        'path': image_path,
                        'original_path': file_info.filename
                    })
    except Exception as e:
        print(f"  Ошибка при извлечении изображений: {e}")
    
    return images

def copy_paragraph_with_images(para, target_doc, images_dict):
    """Копировать параграф с изображениями"""
    new_para = target_doc.add_paragraph()
    
    # Копируем стиль
    try:
        if para.style and para.style.name:
            new_para.style = para.style.name
    except:
        pass
    
    # Копируем выравнивание
    if para.alignment:
        new_para.alignment = para.alignment
    
    # Копируем отступы
    if para.paragraph_format:
        try:
            new_para.paragraph_format.left_indent = para.paragraph_format.left_indent
            new_para.paragraph_format.right_indent = para.paragraph_format.right_indent
            new_para.paragraph_format.first_line_indent = para.paragraph_format.first_line_indent
            new_para.paragraph_format.space_before = para.paragraph_format.space_before
            new_para.paragraph_format.space_after = para.paragraph_format.space_after
        except:
            pass
    
    # Проверяем, есть ли в параграфе изображения
    has_images = False
    
    # Копируем runs с форматированием
    for run in para.runs:
        # Проверяем, содержит ли run изображение
        if run._element.xpath('.//pic:pic'):
            has_images = True
            continue
            
        new_run = new_para.add_run(run.text)
        
        if run.font:
            try:
                if run.font.name:
                    new_run.font.name = run.font.name
                if run.font.size:
                    new_run.font.size = run.font.size
                new_run.font.bold = run.font.bold
                new_run.font.italic = run.font.italic
                if run.font.color and run.font.color.rgb:
                    new_run.font.color.rgb = run.font.color.rgb
            except:
                pass
    
    return has_images

def copy_content_with_images(source_doc, target_doc, images):
    """Копировать содержимое с изображениями"""
    
    # Копируем параграфы
    for para in source_doc.paragraphs:
        if not para.text.strip() and not para._element.xpath('.//pic:pic'):
            continue
            
        # Копируем параграф
        has_images = copy_paragraph_with_images(para, target_doc, {})
        
        # Если в параграфе были изображения, добавляем их
        if has_images and images:
            for image in images:
                try:
                    # Добавляем изображение в документ
                    target_doc.add_picture(image['path'], width=Inches(5))
                    # Добавляем подпись к изображению
                    caption = target_doc.add_paragraph(f"Рисунок: {image['name']}")
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption.style = 'Caption'
                except Exception as e:
                    print(f"    Ошибка при добавлении изображения {image['name']}: {e}")
    
    # Копируем таблицы
    for table in source_doc.tables:
        new_table = target_doc.add_table(
            rows=len(table.rows),
            cols=len(table.columns)
        )
        
        try:
            if table.style:
                new_table.style = table.style.name
        except:
            pass
        
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                new_cell = new_table.cell(i, j)
                
                for cell_para in cell.paragraphs:
                    if cell_para.text.strip():
                        cell_new_para = new_cell.add_paragraph()
                        
                        for run in cell_para.runs:
                            cell_new_run = cell_new_para.add_run(run.text)
                            
                            if run.font:
                                try:
                                    if run.font.name:
                                        cell_new_run.font.name = run.font.name
                                    if run.font.size:
                                        cell_new_run.font.size = run.font.size
                                    cell_new_run.font.bold = run.font.bold
                                    cell_new_run.font.italic = run.font.italic
                                except:
                                    pass

def simple_copy_with_images(source_doc, target_doc, temp_dir, file_name):
    """Простое копирование с попыткой сохранения изображений"""
    print(f"  Копирование содержимого...")
    
    # Сначала пробуем извлечь изображения
    images = extract_images_from_docx(file_name, temp_dir)
    
    if images:
        print(f"  Найдено изображений: {len(images)}")
        for img in images:
            print(f"    - {img['name']}")
    
    # Копируем все параграфы
    for para in source_doc.paragraphs:
        if not para.text.strip():
            # Проверяем, есть ли в пустом параграфе изображения
            if para._element.xpath('.//a:blip'):
                # Пытаемся добавить изображения
                if images:
                    for image in images:
                        try:
                            target_doc.add_picture(image['path'], width=Inches(5))
                            caption = target_doc.add_paragraph(f"Рисунок: {os.path.basename(image['name'])}")
                            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        except:
                            target_doc.add_paragraph(f"[Изображение: {image['name']}]")
                continue
            
            # Пропускаем полностью пустые параграфы
            continue
            
        new_para = target_doc.add_paragraph()
        
        # Копируем текст и форматирование
        for run in para.runs:
            new_run = new_para.add_run(run.text)
            
            if run.font:
                try:
                    if run.font.name:
                        new_run.font.name = run.font.name
                    if run.font.size:
                        new_run.font.size = run.font.size
                    new_run.font.bold = run.font.bold
                    new_run.font.italic = run.font.italic
                except:
                    pass
    
    # Копируем таблицы
    for table in source_doc.tables:
        new_table = target_doc.add_table(
            rows=len(table.rows),
            cols=len(table.columns)
        )
        
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                new_cell = new_table.cell(i, j)
                new_cell.text = cell.text

def main():
    """Основная функция"""
    print("=" * 60)
    print("ОБЪЕДИНЕНИЕ 4 ОТЧЕТОВ С ИЗОБРАЖЕНИЯМИ")
    print("=" * 60)
    
    # Получаем список файлов
    selected_files = get_selected_files()
    
    # Проверяем существование файлов
    existing_files = []
    for file_name in selected_files:
        if os.path.exists(file_name):
            existing_files.append(file_name)
        else:
            print(f"Файл не найден: {file_name}")
    
    if not existing_files:
        print("Нет файлов для обработки")
        return
    
    print("\nВЫБРАННЫЕ ФАЙЛЫ:")
    for i, file_name in enumerate(existing_files, 1):
        print(f"{i}. {file_name}")
    
    print(f"\nВсего файлов: {len(existing_files)}")
    print("=" * 60)
    
    # Создаем временную директорию для изображений
    temp_dir = tempfile.mkdtemp()
    print(f"Временная директория для изображений: {temp_dir}")
    
    try:
        # Создаем новый документ
        combined_doc = Document()
        
        # Настраиваем поля
        for section in combined_doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Добавляем заголовок
        title = combined_doc.add_heading('ОБЪЕДИНЕННЫЙ ОТЧЕТ С ИЗОБРАЖЕНИЯМИ', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = combined_doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.add_run(f"4 лабораторные работы с рисунками и графиками")
        subtitle_run.bold = True
        
        # Список файлов
        files_list = combined_doc.add_paragraph()
        files_list.add_run("Содержание:").bold = True
        
        for i, file_name in enumerate(existing_files, 1):
            name = os.path.basename(file_name).replace('.docx', '')
            combined_doc.add_paragraph(f"{i}. {name}")
        
        combined_doc.add_page_break()
        
        # Обрабатываем файлы
        total_images = 0
        for i, file_name in enumerate(existing_files, 1):
            print(f"\nОбработка {i}/{len(existing_files)}: {file_name}")
            
            try:
                source_doc = Document(file_name)
                
                # Заголовок отчета
                report_title = combined_doc.add_heading(f"Лабораторная работа {i}", 1)
                name = os.path.basename(file_name).replace('.docx', '')
                combined_doc.add_heading(name, 2)
                
                # Извлекаем изображения из этого файла
                images = extract_images_from_docx(file_name, temp_dir)
                total_images += len(images)
                
                if images:
                    print(f"  Найдено изображений в файле: {len(images)}")
                
                # Копируем содержимое
                simple_copy_with_images(source_doc, combined_doc, temp_dir, file_name)
                
                # Разрыв страницы (кроме последнего)
                if i < len(existing_files):
                    combined_doc.add_page_break()
                    
                print(f"  Успешно")
                
            except Exception as e:
                print(f"  Ошибка: {e}")
                combined_doc.add_paragraph(f"[Ошибка при обработке файла {file_name}]")
                
                if i < len(existing_files):
                    combined_doc.add_page_break()
        
        # Сохраняем
        output_name = "объединенный_4_отчета_с_картинками.docx"
        combined_doc.save(output_name)
        
        print("\n" + "=" * 60)
        print(f"Создан файл: {output_name}")
        print(f"Размер: {os.path.getsize(output_name)} байт")
        print(f"Всего изображений найдено: {total_images}")
        print("=" * 60)
        
        print("\nГотово! Отчет создан из 4 файлов:")
        for i, file_name in enumerate(existing_files, 1):
            print(f"{i}. {file_name}")
            
    finally:
        # Очищаем временную директорию
        try:
            shutil.rmtree(temp_dir)
            print(f"\nВременная директория очищена: {temp_dir}")
        except:
            pass

if __name__ == "__main__":
    main()