#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Финальный скрипт для объединения .docx файлов с максимальным сохранением форматирования
"""

import os
import glob
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def get_docx_files():
    """Получить список всех .docx файлов в текущей директории"""
    docx_files = []
    
    # Ищем все файлы .docx
    for file_path in glob.glob("*.docx"):
        # Пропускаем файлы, которые могут быть результатом объединения
        if not file_path.startswith("объединенный_отчет") and not file_path.startswith("combined_report"):
            docx_files.append(file_path)
    
    # Сортируем файлы по имени для удобства
    docx_files.sort()
    
    return docx_files

def copy_content_preserving_format(source_doc, target_doc):
    """Копировать содержимое с максимальным сохранением форматирования"""
    
    # Копируем все параграфы
    for para in source_doc.paragraphs:
        if not para.text.strip():  # Пропускаем пустые параграфы
            continue
            
        # Создаем новый параграф
        new_para = target_doc.add_paragraph()
        
        # Пытаемся скопировать стиль, но не падаем если его нет
        try:
            if para.style and para.style.name:
                new_para.style = para.style.name
        except:
            pass  # Игнорируем ошибки со стилями
        
        # Копируем выравнивание
        if para.alignment:
            new_para.alignment = para.alignment
        
        # Копируем отступы
        if para.paragraph_format:
            try:
                new_para.paragraph_format.left_indent = para.paragraph_format.left_indent
                new_para.paragraph_format.right_indent = para.paragraph_format.right_indent
                new_para.paragraph_format.first_line_indent = para.paragraph_format.first_line_indent
                new_para.paragraph_format.space_before = para.paragraph_format.space_before
                new_para.paragraph_format.space_after = para.paragraph_format.space_after
                new_para.paragraph_format.line_spacing = para.paragraph_format.line_spacing
            except:
                pass
        
        # Копируем все runs с их форматированием
        for run in para.runs:
            new_run = new_para.add_run(run.text)
            
            # Копируем свойства шрифта
            if run.font:
                try:
                    if run.font.name:
                        new_run.font.name = run.font.name
                    if run.font.size:
                        new_run.font.size = run.font.size
                    new_run.font.bold = run.font.bold
                    new_run.font.italic = run.font.italic
                    new_run.font.underline = run.font.underline
                    if run.font.color and run.font.color.rgb:
                        new_run.font.color.rgb = run.font.color.rgb
                except:
                    pass
    
    # Копируем все таблицы
    for table in source_doc.tables:
        # Создаем таблицу с теми же размерами
        new_table = target_doc.add_table(
            rows=len(table.rows),
            cols=len(table.columns)
        )
        
        # Пытаемся скопировать стиль таблицы
        try:
            if table.style:
                new_table.style = table.style.name
        except:
            pass
        
        # Копируем содержимое ячеек
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                new_cell = new_table.cell(i, j)
                
                # Копируем текст из ячейки
                for cell_para in cell.paragraphs:
                    if cell_para.text.strip():
                        cell_new_para = new_cell.add_paragraph()
                        
                        # Копируем runs из ячейки
                        for run in cell_para.runs:
                            cell_new_run = cell_new_para.add_run(run.text)
                            
                            # Копируем форматирование
                            if run.font:
                                try:
                                    if run.font.name:
                                        cell_new_run.font.name = run.font.name
                                    if run.font.size:
                                        cell_new_run.font.size = run.font.size
                                    cell_new_run.font.bold = run.font.bold
                                    cell_new_run.font.italic = run.font.italic
                                    if run.font.color and run.font.color.rgb:
                                        cell_new_run.font.color.rgb = run.font.color.rgb
                                except:
                                    pass

def create_final_combined_report(docx_files):
    """Создать финальный объединенный отчет"""
    if not docx_files:
        print("Не найдено .docx файлов для объединения")
        return
    
    print(f"Найдено {len(docx_files)} .docx файлов:")
    for i, file_path in enumerate(docx_files, 1):
        print(f"  {i}. {file_path}")
    
    # Создаем новый документ
    combined_doc = Document()
    
    # Настраиваем поля страницы
    sections = combined_doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Добавляем простой заголовок
    title = combined_doc.add_heading('ОБЪЕДИНЕННЫЙ ОТЧЕТ ПО ЛАБОРАТОРНЫМ РАБОТАМ', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = combined_doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(f"Всего отчетов: {len(docx_files)}")
    subtitle_run.bold = True
    
    combined_doc.add_page_break()
    
    # Обрабатываем каждый файл
    for i, file_path in enumerate(docx_files, 1):
        print(f"Обработка файла {i}/{len(docx_files)}: {file_path}")
        
        try:
            # Открываем исходный документ
            source_doc = Document(file_path)
            
            # Добавляем разделитель с названием файла
            file_title = combined_doc.add_heading(f"ОТЧЕТ {i}: {os.path.basename(file_path).replace('.docx', '')}", 1)
            
            # Копируем содержимое с сохранением форматирования
            copy_content_preserving_format(source_doc, combined_doc)
            
            # Добавляем разрыв страницы между отчетами (кроме последнего)
            if i < len(docx_files):
                combined_doc.add_page_break()
                
        except Exception as e:
            print(f"  Ошибка при обработке файла: {e}")
            # Добавляем сообщение об ошибке
            error_para = combined_doc.add_paragraph(f"[Файл {os.path.basename(file_path)} содержит ошибки форматирования]")
            error_run = error_para.add_run(" - часть содержимого может отображаться некорректно")
            error_run.italic = True
            
            if i < len(docx_files):
                combined_doc.add_page_break()
    
    # Сохраняем объединенный документ
    output_filename = "объединенный_отчет_финальный.docx"
    combined_doc.save(output_filename)
    
    print(f"\nОбъединенный отчет сохранен в файл: {output_filename}")
    print(f"Размер файла: {os.path.getsize(output_filename)} байт")
    
    return output_filename

def main():
    """Основная функция"""
    print("=" * 60)
    print("Финальный скрипт для объединения .docx отчетов")
    print("=" * 60)
    
    # Получаем список .docx файлов
    docx_files = get_docx_files()
    
    if not docx_files:
        print("В текущей директории не найдено .docx файлов")
        return
    
    # Создаем финальный объединенный документ
    output_file = create_final_combined_report(docx_files)
    
    if output_file:
        print("\n" + "=" * 60)
        print("Готово! Объединенный отчет создан успешно.")
        print("=" * 60)
        print(f"Файл: {os.path.abspath(output_file)}")
        print("\nОсобенности:")
        print("1. Сохранены исходные шрифты и стили из каждого отчета")
        print("2. Сохранено форматирование текста (жирный, курсив, цвет)")
        print("3. Сохранены таблицы с их содержимым")
        print("4. Добавлены заголовки для каждого отчета")
        print("5. Между отчетами добавлены разрывы страниц")

if __name__ == "__main__":
    main()