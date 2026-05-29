#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Скрипт для объединения 4 выбранных .docx файлов с сохранением форматирования
"""

import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def get_selected_files():
    """Получить список выбранных файлов"""
    return [
        "Kurnosov_lab1-2_rep.docx",
        "Kurnosov_lab3-4_rep.docx", 
        "Kurnosov_lab5_report.docx",
        "Kurnosov_lab6_report.docx"
    ]

def copy_content_with_formatting(source_doc, target_doc):
    """Копировать содержимое с сохранением форматирования"""
    
    # Копируем параграфы
    for para in source_doc.paragraphs:
        if not para.text.strip():
            continue
            
        new_para = target_doc.add_paragraph()
        
        # Копируем стиль
        try:
            if para.style and para.style.name:
                new_para.style = para.style.name
        except:
            pass
        
        # Копируем выравнивание
        if para.alignment:
            new_para.alignment = para.alignment
        
        # Копируем отступы
        if para.paragraph_format:
            try:
                new_para.paragraph_format.left_indent = para.paragraph_format.left_indent
                new_para.paragraph_format.right_indent = para.paragraph_format.right_indent
                new_para.paragraph_format.first_line_indent = para.paragraph_format.first_line_indent
                new_para.paragraph_format.space_before = para.paragraph_format.space_before
                new_para.paragraph_format.space_after = para.paragraph_format.space_after
            except:
                pass
        
        # Копируем runs с форматированием
        for run in para.runs:
            new_run = new_para.add_run(run.text)
            
            if run.font:
                try:
                    if run.font.name:
                        new_run.font.name = run.font.name
                    if run.font.size:
                        new_run.font.size = run.font.size
                    new_run.font.bold = run.font.bold
                    new_run.font.italic = run.font.italic
                    if run.font.color and run.font.color.rgb:
                        new_run.font.color.rgb = run.font.color.rgb
                except:
                    pass
    
    # Копируем таблицы
    for table in source_doc.tables:
        new_table = target_doc.add_table(
            rows=len(table.rows),
            cols=len(table.columns)
        )
        
        try:
            if table.style:
                new_table.style = table.style.name
        except:
            pass
        
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                new_cell = new_table.cell(i, j)
                
                for cell_para in cell.paragraphs:
                    if cell_para.text.strip():
                        cell_new_para = new_cell.add_paragraph()
                        
                        for run in cell_para.runs:
                            cell_new_run = cell_new_para.add_run(run.text)
                            
                            if run.font:
                                try:
                                    if run.font.name:
                                        cell_new_run.font.name = run.font.name
                                    if run.font.size:
                                        cell_new_run.font.size = run.font.size
                                    cell_new_run.font.bold = run.font.bold
                                    cell_new_run.font.italic = run.font.italic
                                except:
                                    pass

def main():
    """Основная функция"""
    print("=" * 60)
    print("ОБЪЕДИНЕНИЕ 4 ВЫБРАННЫХ ОТЧЕТОВ")
    print("=" * 60)
    
    # Получаем список файлов
    selected_files = get_selected_files()
    
    # Проверяем существование файлов
    existing_files = []
    for file_name in selected_files:
        if os.path.exists(file_name):
            existing_files.append(file_name)
        else:
            print(f"Файл не найден: {file_name}")
    
    if not existing_files:
        print("Нет файлов для обработки")
        return
    
    print("\nВЫБРАННЫЕ ФАЙЛЫ:")
    for i, file_name in enumerate(existing_files, 1):
        print(f"{i}. {file_name}")
    
    print(f"\nВсего файлов: {len(existing_files)}")
    print("=" * 60)
    
    # Создаем новый документ
    combined_doc = Document()
    
    # Настраиваем поля
    for section in combined_doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Добавляем заголовок
    title = combined_doc.add_heading('ОБЪЕДИНЕННЫЙ ОТЧЕТ', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = combined_doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(f"4 лабораторные работы")
    subtitle_run.bold = True
    
    # Список файлов
    files_list = combined_doc.add_paragraph()
    files_list.add_run("Содержание:").bold = True
    
    for i, file_name in enumerate(existing_files, 1):
        name = os.path.basename(file_name).replace('.docx', '')
        combined_doc.add_paragraph(f"{i}. {name}")
    
    combined_doc.add_page_break()
    
    # Обрабатываем файлы
    for i, file_name in enumerate(existing_files, 1):
        print(f"Обработка {i}/{len(existing_files)}: {file_name}")
        
        try:
            source_doc = Document(file_name)
            
            # Заголовок отчета
            report_title = combined_doc.add_heading(f"Лабораторная работа {i}", 1)
            name = os.path.basename(file_name).replace('.docx', '')
            combined_doc.add_heading(name, 2)
            
            # Копируем содержимое
            copy_content_with_formatting(source_doc, combined_doc)
            
            # Разрыв страницы (кроме последнего)
            if i < len(existing_files):
                combined_doc.add_page_break()
                
            print(f"  Успешно")
            
        except Exception as e:
            print(f"  Ошибка: {e}")
            combined_doc.add_paragraph(f"[Ошибка при обработке файла {file_name}]")
            
            if i < len(existing_files):
                combined_doc.add_page_break()
    
    # Сохраняем
    output_name = "объединенный_4_отчета.docx"
    combined_doc.save(output_name)
    
    print("\n" + "=" * 60)
    print(f"Создан файл: {output_name}")
    print(f"Размер: {os.path.getsize(output_name)} байт")
    print("=" * 60)
    
    print("\nГотово! Отчет создан из 4 файлов:")
    for i, file_name in enumerate(existing_files, 1):
        print(f"{i}. {file_name}")

if __name__ == "__main__":
    main()