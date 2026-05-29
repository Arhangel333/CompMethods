#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Простой скрипт для объединения 4 отчетов с максимальным сохранением содержимого
"""

import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def get_files():
    """Получить список файлов"""
    return [
        "Kurnosov_lab1-2_rep.docx",
        "Kurnosov_lab3-4_rep.docx", 
        "Kurnosov_lab5_report.docx",
        "Kurnosov_lab6_report.docx"
    ]

def add_document_content(source_path, target_doc, doc_number):
    """Добавить содержимое документа в целевой документ"""
    print(f"  Чтение файла...")
    
    try:
        source_doc = Document(source_path)
        
        # Добавляем заголовок
        name = os.path.basename(source_path).replace('.docx', '')
        target_doc.add_heading(f"Лабораторная работа {doc_number}: {name}", 1)
        
        # Просто копируем все параграфы
        para_count = 0
        for para in source_doc.paragraphs:
            if para.text.strip() or para._element.xpath('.//w:drawing'):
                new_para = target_doc.add_paragraph()
                
                # Копируем выравнивание
                if para.alignment:
                    new_para.alignment = para.alignment
                
                # Копируем текст и форматирование
                for run in para.runs:
                    new_run = new_para.add_run(run.text)
                    
                    # Сохраняем форматирование шрифта
                    if run.font:
                        try:
                            if run.font.name:
                                new_run.font.name = run.font.name
                            if run.font.size:
                                new_run.font.size = run.font.size
                            new_run.font.bold = run.font.bold
                            new_run.font.italic = run.font.italic
                        except:
                            pass
                
                para_count += 1
        
        # Копируем таблицы
        table_count = 0
        for table in source_doc.tables:
            # Создаем таблицу
            new_table = target_doc.add_table(
                rows=len(table.rows),
                cols=len(table.columns)
            )
            
            # Копируем содержимое ячеек
            for i, row in enumerate(table.rows):
                for j, cell in enumerate(row.cells):
                    new_cell = new_table.cell(i, j)
                    new_cell.text = cell.text
            
            table_count += 1
        
        print(f"  Скопировано: {para_count} параграфов, {table_count} таблиц")
        return True
        
    except Exception as e:
        print(f"  Ошибка: {e}")
        target_doc.add_paragraph(f"[Ошибка при обработке файла {os.path.basename(source_path)}]")
        return False

def main():
    """Основная функция"""
    print("=" * 60)
    print("ПРОСТОЕ ОБЪЕДИНЕНИЕ 4 ОТЧЕТОВ")
    print("=" * 60)
    
    files = get_files()
    
    # Проверяем файлы
    existing_files = []
    for f in files:
        if os.path.exists(f):
            existing_files.append(f)
        else:
            print(f"Файл не найден: {f}")
    
    if not existing_files:
        print("Нет файлов для обработки")
        return
    
    print("\nФАЙЛЫ ДЛЯ ОБЪЕДИНЕНИЯ:")
    for i, f in enumerate(existing_files, 1):
        print(f"{i}. {f}")
    
    print(f"\nВсего: {len(existing_files)} файлов")
    print("=" * 60)
    
    # Создаем новый документ
    doc = Document()
    
    # Настраиваем поля
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Заголовок
    title = doc.add_heading('ОБЪЕДИНЕННЫЙ ОТЧЕТ', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f"4 лабораторные работы").alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Содержание
    doc.add_page_break()
    doc.add_heading("СОДЕРЖАНИЕ", 1)
    
    for i, f in enumerate(existing_files, 1):
        name = os.path.basename(f).replace('.docx', '')
        doc.add_paragraph(f"{i}. {name}")
    
    doc.add_page_break()
    
    # Обрабатываем файлы
    for i, f in enumerate(existing_files, 1):
        print(f"\nФайл {i}/{len(existing_files)}: {os.path.basename(f)}")
        add_document_content(f, doc, i)
        
        if i < len(existing_files):
            doc.add_page_break()
    
    # Сохраняем
    output = "объединенный_4_отчета_полный.docx"
    doc.save(output)
    
    print("\n" + "=" * 60)
    print(f"Создан: {output}")
    print(f"Размер: {os.path.getsize(output):,} байт")
    print("=" * 60)
    
    print("\nВКЛЮЧЕНЫ:")
    for i, f in enumerate(existing_files, 1):
        print(f"{i}. {os.path.basename(f)}")

if __name__ == "__main__":
    main()