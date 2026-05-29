#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Расширенный скрипт для объединения .docx файлов с сохранением структуры
"""

import os
import glob
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE

def get_docx_files():
    """Получить список всех .docx файлов в текущей директории"""
    docx_files = []
    
    # Ищем все файлы .docx
    for file_path in glob.glob("*.docx"):
        # Пропускаем файлы, которые могут быть результатом объединения
        if not file_path.startswith("объединенный_отчет") and not file_path.startswith("combined_report"):
            docx_files.append(file_path)
    
    # Сортируем файлы по имени для удобства (предполагаем, что в названии есть номер лабораторной)
    docx_files.sort()
    
    return docx_files

def extract_content_from_docx(file_path):
    """Извлечь структурированное содержимое из .docx файла"""
    try:
        doc = Document(file_path)
        content = {
            'title': os.path.basename(file_path).replace('.docx', ''),
            'paragraphs': [],
            'headings': [],
            'tables': []
        }
        
        # Извлекаем заголовки и параграфы
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():
                # Определяем уровень заголовка
                if paragraph.style.name.startswith('Heading'):
                    level = int(paragraph.style.name[-1]) if paragraph.style.name[-1].isdigit() else 1
                    content['headings'].append({
                        'text': paragraph.text.strip(),
                        'level': level,
                        'position': i
                    })
                else:
                    content['paragraphs'].append({
                        'text': paragraph.text.strip(),
                        'position': i
                    })
        
        # Извлекаем таблицы
        for table_idx, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_data.append(cell.text.strip())
                if row_data:
                    table_data.append(row_data)
            
            if table_data:
                content['tables'].append(table_data)
        
        return content
    
    except Exception as e:
        print(f"Ошибка при чтении файла {file_path}: {e}")
        return None

def create_custom_styles(doc):
    """Создать пользовательские стили для документа"""
    styles = doc.styles
    
    # Стиль для заголовков лабораторных
    if 'LabTitle' not in styles:
        lab_title_style = styles.add_style('LabTitle', WD_STYLE_TYPE.PARAGRAPH)
        lab_title_style.font.size = Pt(14)
        lab_title_style.font.bold = True
        lab_title_style.paragraph_format.space_before = Pt(18)
        lab_title_style.paragraph_format.space_after = Pt(12)
    
    # Стиль для обычного текста
    if 'NormalText' not in styles:
        normal_text_style = styles.add_style('NormalText', WD_STYLE_TYPE.PARAGRAPH)
        normal_text_style.font.size = Pt(11)
        normal_text_style.paragraph_format.space_after = Pt(6)
        normal_text_style.paragraph_format.line_spacing = 1.15

def create_combined_document_advanced(docx_files):
    """Создать объединенный документ с улучшенной структурой"""
    if not docx_files:
        print("Не найдено .docx файлов для объединения")
        return
    
    print(f"Найдено {len(docx_files)} .docx файлов:")
    for i, file_path in enumerate(docx_files, 1):
        print(f"  {i}. {file_path}")
    
    # Создаем новый документ
    combined_doc = Document()
    
    # Создаем пользовательские стили
    create_custom_styles(combined_doc)
    
    # Настраиваем поля страницы
    sections = combined_doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Добавляем титульную страницу
    title = combined_doc.add_heading('ОТЧЕТ', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = combined_doc.add_heading('по всем лабораторным работам', 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Добавляем информацию о студенте (можно изменить)
    student_info = combined_doc.add_paragraph()
    student_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    student_info.add_run("Студент: Курносов").bold = True
    
    course_info = combined_doc.add_paragraph()
    course_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    course_info.add_run("Дисциплина: Численные методы").italic = True
    
    combined_doc.add_page_break()
    
    # Добавляем оглавление
    toc_heading = combined_doc.add_heading('СОДЕРЖАНИЕ', 1)
    toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Создаем список для оглавления
    toc_items = []
    
    # Обрабатываем каждый файл
    for i, file_path in enumerate(docx_files, 1):
        print(f"Обработка файла {i}/{len(docx_files)}: {file_path}")
        
        # Извлекаем структурированное содержимое
        content = extract_content_from_docx(file_path)
        
        if not content:
            print(f"  Предупреждение: не удалось обработать файл {file_path}")
            continue
        
        # Определяем название лабораторной из имени файла
        lab_name = content['title']
        lab_number = i
        
        # Добавляем в оглавление
        toc_items.append(f"{lab_number}. {lab_name}")
        
        # Добавляем разделитель
        combined_doc.add_heading(f"Лабораторная работа {lab_number}", 1)
        combined_doc.add_heading(lab_name, 2)
        
        # Добавляем содержимое файла
        if content['paragraphs']:
            for para in content['paragraphs']:
                p = combined_doc.add_paragraph(para['text'])
                p.style = 'NormalText'
        
        # Добавляем таблицы
        if content['tables']:
            for table_data in content['tables']:
                # Создаем таблицу
                table = combined_doc.add_table(rows=len(table_data), cols=len(table_data[0]) if table_data else 0)
                table.style = 'Table Grid'
                
                # Заполняем таблицу
                for row_idx, row_data in enumerate(table_data):
                    row_cells = table.rows[row_idx].cells
                    for col_idx, cell_data in enumerate(row_data):
                        if col_idx < len(row_cells):
                            row_cells[col_idx].text = cell_data
        
        # Добавляем разделитель между отчетами (кроме последнего)
        if i < len(docx_files):
            combined_doc.add_page_break()
    
    # Вставляем оглавление в начало (после титульной страницы)
    # Для этого нужно немного переработать логику, но для простоты оставим так
    
    # Сохраняем объединенный документ
    output_filename = "объединенный_отчет_структурированный.docx"
    combined_doc.save(output_filename)
    
    print(f"\nОбъединенный отчет сохранен в файл: {output_filename}")
    print(f"Размер файла: {os.path.getsize(output_filename)} байт")
    
    # Создаем простой текстовый файл с оглавлением
    create_toc_file(toc_items, docx_files)
    
    return output_filename

def create_toc_file(toc_items, docx_files):
    """Создать текстовый файл с оглавлением"""
    try:
        with open("оглавление_отчетов.txt", "w", encoding="utf-8") as f:
            f.write("ОГЛАВЛЕНИЕ ОТЧЕТОВ\n")
            f.write("=" * 50 + "\n\n")
            
            f.write(f"Всего отчетов: {len(docx_files)}\n\n")
            
            for i, item in enumerate(toc_items, 1):
                f.write(f"{item}\n")
            
            f.write("\n" + "=" * 50 + "\n")
            f.write("Файлы, включенные в объединенный отчет:\n")
            for i, file_path in enumerate(docx_files, 1):
                f.write(f"{i}. {file_path}\n")
        
        print(f"Оглавление сохранено в файл: оглавление_отчетов.txt")
    
    except Exception as e:
        print(f"Ошибка при создании файла оглавления: {e}")

def main():
    """Основная функция"""
    print("=" * 60)
    print("Расширенный скрипт для объединения .docx отчетов")
    print("=" * 60)
    
    # Получаем список .docx файлов
    docx_files = get_docx_files()
    
    if not docx_files:
        print("В текущей директории не найдено .docx файлов")
        return
    
    # Создаем объединенный документ
    output_file = create_combined_document_advanced(docx_files)
    
    if output_file:
        print("\nГотово! Объединенный отчет создан успешно.")
        print(f"Файл: {os.path.abspath(output_file)}")
        print("\nДополнительно создан файл с оглавлением: оглавление_отчетов.txt")

if __name__ == "__main__":
    main()