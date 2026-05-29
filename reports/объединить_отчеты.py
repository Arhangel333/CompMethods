#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Скрипт для объединения текста из нескольких .docx файлов в один документ
"""

import os
import glob
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def get_docx_files():
    """Получить список всех .docx файлов в текущей директории"""
    docx_files = []
    
    # Ищем все файлы .docx
    for file_path in glob.glob("*.docx"):
        # Пропускаем файлы, которые могут быть результатом объединения
        if not file_path.startswith("объединенный_отчет") and not file_path.startswith("combined_report"):
            docx_files.append(file_path)
    
    # Сортируем файлы по имени для удобства
    docx_files.sort()
    
    return docx_files

def extract_text_from_docx(file_path):
    """Извлечь текст из .docx файла"""
    try:
        doc = Document(file_path)
        text_content = []
        
        # Извлекаем текст из всех параграфов
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # Пропускаем пустые параграфы
                text_content.append(paragraph.text)
        
        # Извлекаем текст из таблиц
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text)
                if row_text:
                    text_content.append(" | ".join(row_text))
        
        return "\n".join(text_content), os.path.basename(file_path)
    
    except Exception as e:
        print(f"Ошибка при чтении файла {file_path}: {e}")
        return "", os.path.basename(file_path)

def create_combined_document(docx_files):
    """Создать объединенный документ из всех .docx файлов"""
    if not docx_files:
        print("Не найдено .docx файлов для объединения")
        return
    
    print(f"Найдено {len(docx_files)} .docx файлов:")
    for i, file_path in enumerate(docx_files, 1):
        print(f"  {i}. {file_path}")
    
    # Создаем новый документ
    combined_doc = Document()
    
    # Добавляем заголовок
    title = combined_doc.add_heading('Объединенный отчет по всем лабораторным работам', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Добавляем информацию о файлах
    info_para = combined_doc.add_paragraph()
    info_para.add_run(f"Объединено {len(docx_files)} отчетов:").bold = True
    
    for i, file_path in enumerate(docx_files, 1):
        info_para = combined_doc.add_paragraph(f"{i}. {file_path}")
    
    combined_doc.add_paragraph()  # Пустая строка
    
    # Обрабатываем каждый файл
    for i, file_path in enumerate(docx_files, 1):
        print(f"Обработка файла {i}/{len(docx_files)}: {file_path}")
        
        # Добавляем разделитель
        combined_doc.add_heading(f"Отчет {i}: {os.path.basename(file_path)}", 1)
        
        # Извлекаем текст из файла
        text_content, filename = extract_text_from_docx(file_path)
        
        if text_content:
            # Добавляем текст в документ
            para = combined_doc.add_paragraph(text_content)
            
            # Настраиваем шрифт (опционально)
            for run in para.runs:
                run.font.size = Pt(11)
        else:
            combined_doc.add_paragraph(f"Не удалось извлечь текст из файла {filename}")
        
        # Добавляем разделитель между отчетами
        if i < len(docx_files):
            combined_doc.add_page_break()
    
    # Сохраняем объединенный документ
    output_filename = "объединенный_отчет_все_лабораторные.docx"
    combined_doc.save(output_filename)
    
    print(f"\nОбъединенный отчет сохранен в файл: {output_filename}")
    print(f"Размер файла: {os.path.getsize(output_filename)} байт")
    
    return output_filename

def main():
    """Основная функция"""
    print("=" * 60)
    print("Скрипт для объединения .docx отчетов")
    print("=" * 60)
    
    # Получаем список .docx файлов
    docx_files = get_docx_files()
    
    if not docx_files:
        print("В текущей директории не найдено .docx файлов")
        return
    
    # Создаем объединенный документ
    output_file = create_combined_document(docx_files)
    
    if output_file:
        print("\nГотово! Объединенный отчет создан успешно.")
        print(f"Файл: {os.path.abspath(output_file)}")

if __name__ == "__main__":
    main()