#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Скрипт для анализа и объединения только уникальных лабораторных работ
"""

import os
import glob
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def get_docx_files():
    """Получить список всех .docx файлов в текущей директории"""
    docx_files = []
    
    # Ищем все файлы .docx
    for file_path in glob.glob("*.docx"):
        # Пропускаем файлы, которые могут быть результатом объединения
        if not file_path.startswith("объединенный_отчет") and not file_path.startswith("combined_report"):
            docx_files.append(file_path)
    
    # Сортируем файлы по имени для удобства
    docx_files.sort()
    
    return docx_files

def analyze_lab_numbers(files):
    """Проанализировать номера лабораторных работ в файлах"""
    print("=" * 60)
    print("АНАЛИЗ ФАЙЛОВ НА ПОВТОРЯЮЩИЕСЯ ЛАБОРАТОРНЫЕ РАБОТЫ")
    print("=" * 60)
    
    lab_patterns = []
    file_analysis = []
    
    for file_path in files:
        filename = os.path.basename(file_path)
        
        # Пытаемся извлечь номера лабораторных из имени файла
        lab_numbers = []
        
        # Ищем паттерны типа "lab1-2", "lab3-4", "lab5", "lab6"
        patterns = [
            r'lab(\d+)-(\d+)',  # lab1-2, lab3-4
            r'lab(\d+)',        # lab5, lab6
            r'лаб(\d+)-(\d+)',  # лаб1-2
            r'лаб(\d+)',        # лаб5
        ]
        
        for pattern in patterns:
            matches = re.findall(pattern, filename, re.IGNORECASE)
            if matches:
                if isinstance(matches[0], tuple):  # lab1-2
                    start, end = matches[0]
                    lab_numbers.extend(range(int(start), int(end) + 1))
                else:  # lab5
                    lab_numbers.append(int(matches[0]))
                break
        
        # Если не нашли в имени файла, попробуем заглянуть в содержимое
        if not lab_numbers:
            try:
                doc = Document(file_path)
                # Берем первые 3 параграфа для анализа
                content_sample = " ".join([para.text for para in doc.paragraphs[:3]])
                
                # Ищем в содержимом
                for pattern in patterns:
                    matches = re.findall(pattern, content_sample, re.IGNORECASE)
                    if matches:
                        if isinstance(matches[0], tuple):  # lab1-2
                            start, end = matches[0]
                            lab_numbers.extend(range(int(start), int(end) + 1))
                        else:  # lab5
                            lab_numbers.append(int(matches[0]))
                        break
            except:
                pass
        
        file_analysis.append({
            'path': file_path,
            'filename': filename,
            'lab_numbers': sorted(lab_numbers) if lab_numbers else [],
            'is_duplicate': False
        })
    
    # Определяем дубликаты
    seen_labs = set()
    for file_info in file_analysis:
        if file_info['lab_numbers']:
            # Проверяем, есть ли пересечение с уже увиденными лабораторными
            current_labs = set(file_info['lab_numbers'])
            if current_labs.intersection(seen_labs):
                file_info['is_duplicate'] = True
            else:
                seen_labs.update(current_labs)
    
    # Выводим анализ
    print("\nАНАЛИЗ ФАЙЛОВ:")
    print("-" * 60)
    
    unique_count = 0
    duplicate_count = 0
    
    for i, file_info in enumerate(file_analysis, 1):
        status = "ДУБЛИКАТ" if file_info['is_duplicate'] else "УНИКАЛЬНЫЙ"
        labs = file_info['lab_numbers'] if file_info['lab_numbers'] else ["не определены"]
        
        print(f"{i:2}. {status:12} {file_info['filename']:35} Лабы: {labs}")
        
        if file_info['is_duplicate']:
            duplicate_count += 1
        else:
            unique_count += 1
    
    print("-" * 60)
    print(f"Всего файлов: {len(files)}")
    print(f"Уникальных лабораторных: {unique_count}")
    print(f"Дубликатов: {duplicate_count}")
    print("=" * 60)
    
    return file_analysis

def select_unique_files(file_analysis):
    """Выбрать только уникальные файлы"""
    unique_files = []
    
    print("\nВЫБРАНЫ ДЛЯ ОБЪЕДИНЕНИЯ:")
    print("-" * 60)
    
    for i, file_info in enumerate(file_analysis, 1):
        if not file_info['is_duplicate']:
            unique_files.append(file_info['path'])
            labs = file_info['lab_numbers'] if file_info['lab_numbers'] else ["не определены"]
            print(f"{i:2}. {file_info['filename']:35} Лабы: {labs}")
    
    print("-" * 60)
    print(f"Всего выбрано: {len(unique_files)} файлов")
    
    return unique_files

def copy_content_preserving_format(source_doc, target_doc):
    """Копировать содержимое с максимальным сохранением форматирования"""
    
    # Копируем все параграфы
    for para in source_doc.paragraphs:
        if not para.text.strip():  # Пропускаем пустые параграфы
            continue
            
        # Создаем новый параграф
        new_para = target_doc.add_paragraph()
        
        # Пытаемся скопировать стиль, но не падаем если его нет
        try:
            if para.style and para.style.name:
                new_para.style = para.style.name
        except:
            pass  # Игнорируем ошибки со стилями
        
        # Копируем выравнивание
        if para.alignment:
            new_para.alignment = para.alignment
        
        # Копируем отступы
        if para.paragraph_format:
            try:
                new_para.paragraph_format.left_indent = para.paragraph_format.left_indent
                new_para.paragraph_format.right_indent = para.paragraph_format.right_indent
                new_para.paragraph_format.first_line_indent = para.paragraph_format.first_line_indent
                new_para.paragraph_format.space_before = para.paragraph_format.space_before
                new_para.paragraph_format.space_after = para.paragraph_format.space_after
                new_para.paragraph_format.line_spacing = para.paragraph_format.line_spacing
            except:
                pass
        
        # Копируем все runs с их форматированием
        for run in para.runs:
            new_run = new_para.add_run(run.text)
            
            # Копируем свойства шрифта
            if run.font:
                try:
                    if run.font.name:
                        new_run.font.name = run.font.name
                    if run.font.size:
                        new_run.font.size = run.font.size
                    new_run.font.bold = run.font.bold
                    new_run.font.italic = run.font.italic
                    new_run.font.underline = run.font.underline
                    if run.font.color and run.font.color.rgb:
                        new_run.font.color.rgb = run.font.color.rgb
                except:
                    pass
    
    # Копируем все таблицы
    for table in source_doc.tables:
        # Создаем таблицу с теми же размерами
        new_table = target_doc.add_table(
            rows=len(table.rows),
            cols=len(table.columns)
        )
        
        # Пытаемся скопировать стиль таблицы
        try:
            if table.style:
                new_table.style = table.style.name
        except:
            pass
        
        # Копируем содержимое ячеек
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                new_cell = new_table.cell(i, j)
                
                # Копируем текст из ячейки
                for cell_para in cell.paragraphs:
                    if cell_para.text.strip():
                        cell_new_para = new_cell.add_paragraph()
                        
                        # Копируем runs из ячейки
                        for run in cell_para.runs:
                            cell_new_run = cell_new_para.add_run(run.text)
                            
                            # Копируем форматирование
                            if run.font:
                                try:
                                    if run.font.name:
                                        cell_new_run.font.name = run.font.name
                                    if run.font.size:
                                        cell_new_run.font.size = run.font.size
                                    cell_new_run.font.bold = run.font.bold
                                    cell_new_run.font.italic = run.font.italic
                                    if run.font.color and run.font.color.rgb:
                                        cell_new_run.font.color.rgb = run.font.color.rgb
                                except:
                                    pass

def create_combined_report(unique_files):
    """Создать объединенный отчет из уникальных файлов"""
    if not unique_files:
        print("Нет уникальных файлов для объединения")
        return
    
    print("\n" + "=" * 60)
    print("СОЗДАНИЕ ОБЪЕДИНЕННОГО ОТЧЕТА ИЗ УНИКАЛЬНЫХ ФАЙЛОВ")
    print("=" * 60)
    
    # Создаем новый документ
    combined_doc = Document()
    
    # Настраиваем поля страницы
    sections = combined_doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Добавляем заголовок
    title = combined_doc.add_heading('ОБЪЕДИНЕННЫЙ ОТЧЕТ ПО УНИКАЛЬНЫМ ЛАБОРАТОРНЫМ РАБОТАМ', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = combined_doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(f"Уникальных отчетов: {len(unique_files)} (дубликаты исключены)")
    subtitle_run.bold = True
    
    combined_doc.add_page_break()
    
    # Обрабатываем каждый уникальный файл
    for i, file_path in enumerate(unique_files, 1):
        filename = os.path.basename(file_path)
        print(f"Обработка файла {i}/{len(unique_files)}: {filename}")
        
        try:
            # Открываем исходный документ
            source_doc = Document(file_path)
            
            # Добавляем разделитель с названием файла
            file_title = combined_doc.add_heading(f"ОТЧЕТ {i}: {filename.replace('.docx', '')}", 1)
            
            # Копируем содержимое с сохранением форматирования
            copy_content_preserving_format(source_doc, combined_doc)
            
            # Добавляем разрыв страницы между отчетами (кроме последнего)
            if i < len(unique_files):
                combined_doc.add_page_break()
                
        except Exception as e:
            print(f"  Ошибка при обработке файла: {e}")
            # Добавляем сообщение об ошибке
            error_para = combined_doc.add_paragraph(f"[Файл {filename} содержит ошибки форматирования]")
            error_run = error_para.add_run(" - часть содержимого может отображаться некорректно")
            error_run.italic = True
            
            if i < len(unique_files):
                combined_doc.add_page_break()
    
    # Сохраняем объединенный документ
    output_filename = "объединенный_отчет_уникальные_лабы.docx"
    combined_doc.save(output_filename)
    
    print(f"\nОбъединенный отчет сохранен в файл: {output_filename}")
    print(f"Размер файла: {os.path.getsize(output_filename)} байт")
    
    return output_filename

def main():
    """Основная функция"""
    print("=" * 60)
    print("СКРИПТ ДЛЯ ОБЪЕДИНЕНИЯ ТОЛЬКО УНИКАЛЬНЫХ ЛАБОРАТОРНЫХ РАБОТ")
    print("=" * 60)
    
    # Получаем список .docx файлов
    all_files = get_docx_files()
    
    if not all_files:
        print("В текущей директории не найдено .docx файлов")
        return
    
    # Анализируем файлы на дубликаты
    file_analysis = analyze_lab_numbers(all_files)
    
    # Выбираем только уникальные файлы
    unique_files = select_unique_files(file_analysis)
    
    if not unique_files:
        print("\nНет уникальных файлов для объединения!")
        return
    
    # Спрашиваем подтверждение
    print("\n" + "=" * 60)
    response = input("Продолжить объединение выбранных файлов? (да/нет): ").strip().lower()
    
    if response not in ['да', 'д', 'yes', 'y']:
        print("Объединение отменено.")
        return
    
    # Создаем объединенный документ из уникальных файлов
    output_file = create_combined_report(unique_files)
    
    if output_file:
        print("\n" + "=" * 60)
        print("ГОТОВО! ОБЪЕДИНЕННЫЙ ОТЧЕТ СОЗДАН УСПЕШНО.")
        print("=" * 60)
        print(f"Файл: {os.path.abspath(output_file)}")
        print("\nОсобенности:")
        print("1. Включены только уникальные лабораторные работы")
        print("2. Дубликаты исключены из объединения")
        print("3. Сохранены исходные шрифты и стили")
        print("4. Сохранено форматирование текста")
        print("5. Сохранены таблицы с их содержимым")

if __name__ == "__main__":
    main()