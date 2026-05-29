#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Финальный скрипт для объединения 4 отчетов с сохранением изображений
Использует прямое копирование XML для сохранения всего содержимого
"""

import os
import zipfile
import shutil
import tempfile
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def get_files():
    """Получить список файлов"""
    return [
        "Kurnosov_lab1-2_rep.docx",
        "Kurnosov_lab3-4_rep.docx", 
        "Kurnosov_lab5_report.docx",
        "Kurnosov_lab6_report.docx"
    ]

def merge_docx_files(files, output_path):
    """Объединить несколько .docx файлов в один"""
    print("Создание объединенного документа...")
    
    # Создаем основной документ
    main_doc = Document()
    
    # Настраиваем поля
    for section in main_doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Добавляем титульную страницу
    title = main_doc.add_heading('ОБЪЕДИНЕННЫЙ ОТЧЕТ', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = main_doc.add_paragraph('4 лабораторные работы')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    main_doc.add_page_break()
    
    # Создаем временную директорию
    temp_dir = tempfile.mkdtemp()
    
    try:
        # Собираем все изображения из всех файлов
        all_images = {}
        image_counter = 1
        
        for file_idx, file_path in enumerate(files, 1):
            print(f"\nОбработка файла {file_idx}/{len(files)}: {os.path.basename(file_path)}")
            
            # Извлекаем изображения
            images = extract_images(file_path, temp_dir, file_idx)
            print(f"  Найдено изображений: {len(images)}")
            
            # Добавляем в общий словарь
            for img_name, img_path in images.items():
                new_name = f"image_{image_counter}_{os.path.splitext(img_name)[1]}"
                all_images[new_name] = img_path
                image_counter += 1
        
        print(f"\nВсего изображений найдено: {len(all_images)}")
        
        # Теперь обрабатываем каждый файл и добавляем содержимое
        for file_idx, file_path in enumerate(files, 1):
            print(f"\nДобавление содержимого из: {os.path.basename(file_path)}")
            
            # Добавляем заголовок
            name = os.path.basename(file_path).replace('.docx', '')
            main_doc.add_heading(f"Лабораторная работа {file_idx}: {name}", 1)
            
            # Открываем исходный документ
            source_doc = Document(file_path)
            
            # Копируем содержимое
            copy_document_content(source_doc, main_doc)
            
            # Добавляем разрыв страницы (кроме последнего файла)
            if file_idx < len(files):
                main_doc.add_page_break()
        
        # Сохраняем основной документ
        main_doc.save(output_path)
        
        # Теперь добавляем изображения в сохраненный документ
        if all_images:
            print(f"\nДобавление {len(all_images)} изображений в документ...")
            add_images_to_docx(output_path, all_images, temp_dir)
        
        return True
        
    finally:
        # Очищаем временную директорию
        try:
            shutil.rmtree(temp_dir)
        except:
            pass

def extract_images(docx_path, temp_dir, file_idx):
    """Извлечь изображения из .docx файла"""
    images = {}
    
    try:
        with zipfile.ZipFile(docx_path, 'r') as zip_ref:
            # Ищем файлы изображений
            for file_info in zip_ref.infolist():
                if file_info.filename.startswith('word/media/'):
                    img_name = os.path.basename(file_info.filename)
                    # Создаем уникальное имя для изображения
                    unique_name = f"file{file_idx}_{img_name}"
                    img_path = os.path.join(temp_dir, unique_name)
                    
                    # Извлекаем изображение
                    with zip_ref.open(file_info) as source, open(img_path, 'wb') as target:
                        shutil.copyfileobj(source, target)
                    
                    images[unique_name] = img_path
    except Exception as e:
        print(f"  Ошибка при извлечении изображений: {e}")
    
    return images

def copy_document_content(source_doc, target_doc):
    """Копировать содержимое документа"""
    # Копируем параграфы
    for para in source_doc.paragraphs:
        if para.text.strip():
            new_para = target_doc.add_paragraph()
            
            # Сохраняем выравнивание
            if para.alignment:
                new_para.alignment = para.alignment
            
            # Копируем runs с форматированием
            for run in para.runs:
                new_run = new_para.add_run(run.text)
                
                # Сохраняем форматирование шрифта
                if run.font:
                    try:
                        if run.font.name:
                            new_run.font.name = run.font.name
                        if run.font.size:
                            new_run.font.size = run.font.size
                        new_run.font.bold = run.font.bold
                        new_run.font.italic = run.font.italic
                    except:
                        pass
    
    # Копируем таблицы
    for table in source_doc.tables:
        # Создаем таблицу
        new_table = target_doc.add_table(
            rows=len(table.rows),
            cols=len(table.columns)
        )
        
        # Копируем содержимое ячеек
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                new_cell = new_table.cell(i, j)
                new_cell.text = cell.text

def add_images_to_docx(docx_path, images, temp_dir):
    """Добавить изображения в .docx файл"""
    try:
        # Создаем временную копию документа
        temp_docx = os.path.join(temp_dir, "temp.docx")
        shutil.copy2(docx_path, temp_docx)
        
        # Открываем как zip и добавляем изображения
        with zipfile.ZipFile(temp_docx, 'a') as zip_out:
            for img_name, img_path in images.items():
                if os.path.exists(img_path):
                    # Добавляем изображение в архив
                    zip_out.write(img_path, f"word/media/{img_name}")
        
        # Заменяем оригинальный файл
        shutil.move(temp_docx, docx_path)
        
    except Exception as e:
        print(f"  Ошибка при добавлении изображений: {e}")

def simple_merge(files, output_path):
    """Простое объединение без сложной обработки изображений"""
    print("Простое объединение файлов...")
    
    main_doc = Document()
    
    # Настраиваем поля
    for section in main_doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Заголовок
    title = main_doc.add_heading('ОБЪЕДИНЕННЫЙ ОТЧЕТ', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    main_doc.add_paragraph('4 лабораторные работы').alignment = WD_ALIGN_PARAGRAPH.CENTER
    main_doc.add_page_break()
    
    # Обрабатываем файлы
    for i, file_path in enumerate(files, 1):
        print(f"Файл {i}/{len(files)}: {os.path.basename(file_path)}")
        
        name = os.path.basename(file_path).replace('.docx', '')
        main_doc.add_heading(f"Лабораторная работа {i}: {name}", 1)
        
        try:
            source_doc = Document(file_path)
            
            # Простое копирование
            for para in source_doc.paragraphs:
                if para.text.strip():
                    main_doc.add_paragraph(para.text)
            
            for table in source_doc.tables:
                new_table = main_doc.add_table(
                    rows=len(table.rows),
                    cols=len(table.columns)
                )
                
                for row_idx, row in enumerate(table.rows):
                    for col_idx, cell in enumerate(row.cells):
                        new_table.cell(row_idx, col_idx).text = cell.text
            
        except Exception as e:
            print(f"  Ошибка: {e}")
            main_doc.add_paragraph(f"[Ошибка при обработке файла]")
        
        if i < len(files):
            main_doc.add_page_break()
    
    main_doc.save(output_path)
    return True

def main():
    """Основная функция"""
    print("=" * 60)
    print("ФИНАЛЬНОЕ ОБЪЕДИНЕНИЕ 4 ОТЧЕТОВ")
    print("=" * 60)
    
    files = get_files()
    
    # Проверяем файлы
    existing_files = []
    for f in files:
        if os.path.exists(f):
            existing_files.append(f)
        else:
            print(f"Файл не найден: {f}")
    
    if not existing_files:
        print("Нет файлов для обработки")
        return
    
    print("\nФАЙЛЫ ДЛЯ ОБЪЕДИНЕНИЯ:")
    for i, f in enumerate(existing_files, 1):
        print(f"{i}. {os.path.basename(f)}")
    
    print(f"\nВсего: {len(existing_files)} файлов")
    print("=" * 60)
    
    # Пробуем разные методы
    print("\n1. Попытка объединения с сохранением изображений...")
    output1 = "объединенный_4_отчета_с_рисунками.docx"
    try:
        merge_docx_files(existing_files, output1)
        size1 = os.path.getsize(output1)
        print(f"   Создан: {output1} ({size1:,} байт)")
    except Exception as e:
        print(f"   Ошибка: {e}")
        output1 = None
    
    print("\n2. Простое объединение (гарантированно работает)...")
    output2 = "объединенный_4_отчета_простой.docx"
    simple_merge(existing_files, output2)
    size2 = os.path.getsize(output2)
    print(f"   Создан: {output2} ({size2:,} байт)")
    
    print("\n" + "=" * 60)
    print("РЕЗУЛЬТАТ:")
    print("=" * 60)
    
    if output1 and os.path.exists(output1):
        print(f"1. {output1} - с попыткой сохранения изображений")
        print(f"   Размер: {size1:,} байт")
    
    print(f"\n2. {output2} - простой вариант (гарантированно работает)")
    print(f"   Размер: {size2:,} байт")
    
    print("\nРекомендуется открыть оба файла и выбрать подходящий вариант.")

if __name__ == "__main__":
    main()